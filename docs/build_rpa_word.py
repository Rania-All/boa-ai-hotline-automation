# -*- coding: utf-8 -*-
"""Génère le fichier Word RPA-Detail-Robot-Architecture-Etapes.docx"""
import os
try:
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
except ImportError:
    print("Installation de python-docx...")
    os.system("pip install python-docx -q")
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()
style = doc.styles['Normal']
style.font.size = Pt(11)
style.font.name = 'Calibri'

def heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.add_run(text).bold = True
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    if level == 1:
        p.runs[0].font.size = Pt(14)
    elif level == 2:
        p.runs[0].font.size = Pt(12)
    return p

def para(doc, text, bullet=False):
    p = doc.add_paragraph(style=doc.styles['List Bullet'] if bullet else 'Normal')
    if not bullet:
        p.paragraph_format.left_indent = Cm(0)
    p.add_run(text)
    p.paragraph_format.space_after = Pt(4)
    return p

# Titre
t = doc.add_paragraph()
t.add_run('RPA – Architecture et étapes (projet Automatisation Hotline)').bold = True
t.runs[0].font.size = Pt(16)
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()
doc.add_paragraph(
    'Document détaillé pour présenter la partie RPA du projet à l’encadrant pédagogique : '
    'architecture, flux, étapes techniques, sécurité et planning.'
)

# 1. Vue d'ensemble
heading(doc, '1. Vue d’ensemble', 1)
doc.add_paragraph(
    'Le projet Automatisation Hotline (version Python) comporte deux parties principales :'
)
para(doc, 'Partie 1 – Chatbot (FAQ) : Un chatbot s’appuie sur un fichier Excel contenant des paires question / réponse (base FAQ). À partir de la question de l’utilisateur, une classification est effectuée (par exemple via RAG ou moteur de similarité) pour retrouver la réponse adaptée ou pour décider si la demande relève du périmètre N1-RR (éligible à l’automatisation RPA).', bullet=True)
para(doc, 'Partie 2 – RPA : Lorsque la classification indique un cas N1-RR, l’application Python (backend avec RAG) déclenche et suit l’exécution des processus UiPath via l’Orchestrator.', bullet=True)
doc.add_paragraph(
    'Ce document décrit l’architecture et les étapes de la partie RPA.'
)

# 2. Architecture RPA
heading(doc, '2. Architecture RPA', 1)
heading(doc, '2.1 Composants principaux', 2)
table = doc.add_table(rows=4, cols=2)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = 'Composant'
hdr[1].text = 'Rôle'
for c in hdr:
    for p in c.paragraphs:
        for r in p.runs:
            r.bold = True
data = [
    ('Application Python (RAG)', 'Backend : classification N1-RR via RAG/IA, vérification des données, appel des APIs UiPath (client HTTP), suivi des jobs, retour au client.'),
    ('UiPath Orchestrator', 'Plateforme centrale : authentification (OAuth2), envoi des jobs aux robots, suivi du statut.'),
    ('UiPath Robots', 'Exécution des processus RPA (traitement N1-RR, automatisation des tâches hotline).'),
]
for i, (a, b) in enumerate(data, 1):
    table.rows[i].cells[0].text = a
    table.rows[i].cells[1].text = b
doc.add_paragraph()

heading(doc, '2.2 Implémentation technique (backend Python)', 2)
para(doc, 'Configuration : variables d’environnement ou fichier .env (UIPATH_BASE_URL, UIPATH_CLIENT_ID, UIPATH_CLIENT_SECRET, UIPATH_FOLDER_ID, UIPATH_RELEASE_KEY).', bullet=True)
para(doc, 'Client Orchestrator : module requests ou httpx pour OAuth2, POST Start Jobs, GET Jobs pour le statut.', bullet=True)
para(doc, 'RAG : chaîne RAG (LLM + base documentaire) pour la classification N1-RR ; le résultat déclenche ou non l’appel RPA.', bullet=True)
para(doc, 'API REST exposée : endpoints /api/rpa/start et /api/rpa/status/{job_key} (ex. FastAPI/Flask).', bullet=True)

# 3. Étapes
heading(doc, '3. Étapes détaillées du déclenchement UiPath depuis Python', 1)
doc.add_paragraph('Les étapes ci-dessous correspondent au flux réel : de la classification jusqu’au retour au client.')

steps = [
    (1, 'Classification N1-RR par le module IA',
     'En partie 1, le chatbot utilise un fichier Excel (questions / réponses FAQ) pour répondre. Une classification est réalisée : soit une réponse est trouvée dans la FAQ, soit la demande est identifiée comme N1-RR et oriente vers le robot RPA. Seules les demandes N1-RR déclenchent un job RPA.'),
    (2, 'Vérification des données requises',
     'Contrôles côté backend avant l’appel à l’API Start Job (données nécessaires au robot présentes et valides).'),
    (3, 'Authentification auprès d’UiPath Orchestrator (OAuth2)',
     'POST vers https://cloud.uipath.com/identity_/connect/token ; client_id, client_secret, scope (OR.Folders OR.Jobs). Récupération du token ; en-tête Authorization: Bearer <token> pour les requêtes Orchestrator.'),
    (4, 'Appel API Start Job via client HTTP (Python)',
     'POST vers /odata/Jobs/UiPath.Server.Configuration.OData.StartJobs. En-têtes : Authorization, X-UIPATH-OrganizationUnitId, Content-Type application/json. Corps : startInfo (ReleaseKey, Strategy, InputArguments, JobsCount, etc.). Parsing de la réponse pour extraire job_key.'),
    (5, 'Transmission des InputArguments en JSON',
     'Données métier (action, payload) sérialisées avec json.dumps(...) dans le corps de la requête Start Job.'),
    (6, 'Suivi du job via API Get Job Status',
     'GET sur l’API OData des Jobs avec filtre $filter=Key eq job_key. Exposé via GET /api/rpa/status/{job_key}. Polling jusqu’à terminaison ou timeout.'),
    (7, 'Retour automatique au client',
     'Retour immédiat job_key + état initial ; le client peut interroger /api/rpa/status/{job_key} pour le suivi. Évolution possible : retour asynchrone (webhook ou polling) une fois le job terminé.'),
]
for num, title, detail in steps:
    heading(doc, f'Étape {num} : {title}', 2)
    para(doc, detail)

# 4. Sécurité
heading(doc, '4. Sécurité et conformité', 1)
table2 = doc.add_table(rows=5, cols=2)
table2.style = 'Table Grid'
h2 = table2.rows[0].cells
h2[0].text = 'Mesure'
h2[1].text = 'Description'
for c in h2:
    for p in c.paragraphs:
        for r in p.runs:
            r.bold = True
sec = [
    ('Masquage des données sensibles', 'Ex. numéro de carte masqué avant envoi ou affichage (côté métier / logs).'),
    ('Journalisation des actions', 'Appels RPA (démarrage, statut) et erreurs loggés pour traçabilité et audit.'),
    ('Gestion des erreurs et escalade', 'En cas d’échec (token, API, timeout) : gestion d’erreur et escalade automatique (notification, repli manuel).'),
    ('Secrets', 'client_id, client_secret dans variables d’environnement ou fichier .env (non versionné).'),
]
for i, (a, b) in enumerate(sec, 1):
    table2.rows[i].cells[0].text = a
    table2.rows[i].cells[1].text = b
doc.add_paragraph()

# 5. KPIs
heading(doc, '5. KPIs attendus', 1)
para(doc, 'Taux de précision de classification : > 85 % (module RAG/IA pour N1-RR).', bullet=True)
para(doc, 'Taux d’automatisation N1-RR : > 60 % des cas N1-RR traités par RPA.', bullet=True)
para(doc, 'Réduction du temps moyen de traitement : objectif global du projet.', bullet=True)

# 6. Planning
heading(doc, '6. Planning (extrait – partie RPA)', 1)
para(doc, 'Mois 1 : Formation UiPath Academy, conception architecture Python (RAG + API), spécifications fonctionnelles et techniques (dont RPA).', bullet=True)
para(doc, 'Mois 2 : Développement chaîne RAG, API REST (FastAPI/Flask) et moteur de décision ; chatbot + fichier Excel FAQ.', bullet=True)
para(doc, 'Mois 3 (S9–S12) – Intégration UiPath : service d’authentification Orchestrator (OAuth2), déclenchement job via API (Start Job + InputArguments), tests avec robots existants, gestion logs et erreurs. Livrable : Prototype end-to-end fonctionnel.', bullet=True)
para(doc, 'Mois 4 : Tests utilisateurs (hotline), optimisation des performances, documentation technique complète. Livrable : Solution prête pour pilote.', bullet=True)

# 7. Résumé
heading(doc, '7. Résumé pour l’encadrant', 1)
para(doc, 'Architecture : Backend Python (RAG + API) ↔ UiPath Orchestrator (OAuth2) ↔ Robots UiPath ; déclenchement et suivi des jobs via appels HTTP (requests/httpx) aux APIs Orchestrator.', bullet=True)
para(doc, 'Étapes clés : Classification N1-RR (chatbot + Excel FAQ) → Vérification des données → Authentification OAuth2 → Start Job (InputArguments en JSON) → Get Job Status → Retour au client.', bullet=True)
para(doc, 'Sécurité : Masquage des données sensibles, journalisation, gestion des erreurs et escalade ; secrets en configuration.', bullet=True)
para(doc, 'Planning RPA : Intégration UiPath au mois 3 (semaines 9–12), prototype end-to-end puis industrialisation en mois 4.', bullet=True)

doc.add_paragraph()
doc.add_paragraph('Ce document peut être utilisé pour une présentation orale ou une annexe du rapport de stage.')

out = os.path.join(os.path.dirname(__file__), 'RPA-Detail-Robot-Architecture-Etapes.docx')
doc.save(out)
print('Fichier Word créé :', out)
